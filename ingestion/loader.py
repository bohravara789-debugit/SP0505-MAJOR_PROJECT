import os
import tempfile
import pdfplumber
import docx
import pandas as pd
from langchain_core.documents import Document

# Try importing OCR tools (optional dependency handling)
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

def load_document(uploaded_file):
    """
    Loads a document (PDF, DOCX, XLSX, CSV, TXT, MD) and returns a list of Document objects.
    Includes OCR fallback for scanned PDFs using pdf2image and pytesseract.
    """
    if uploaded_file is None:
        return []

    file_extension = uploaded_file.name.split(".")[-1].lower()
    text = ""
    
    # Create a temporary file to handle the upload
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
        temp_file.write(uploaded_file.getvalue())
        temp_path = temp_file.name

    try:
        if file_extension == "pdf":
            # 1. Try standard text extraction
            with pdfplumber.open(temp_path) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            
            # 2. OCR Fallback if text is empty (for scanned PDFs)
            if not text.strip() and OCR_AVAILABLE:
                try:
                    # Note: Poppler must be installed and in PATH for this to work on Windows
                    images = convert_from_path(temp_path)
                    for img in images:
                        text += pytesseract.image_to_string(img) + "\n"
                except Exception as e:
                    print(f"OCR warning (Poppler might be missing): {e}")

        elif file_extension == "docx":
            doc = docx.Document(temp_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        
        elif file_extension in ["xlsx", "xls"]:
            # Excel handling
            df = pd.read_excel(temp_path)
            text = df.to_string(index=False)
            
        elif file_extension == "csv":
            # CSV handling
            df = pd.read_csv(temp_path)
            text = df.to_string(index=False)
            
        elif file_extension in ["txt", "md"]:
            # Text and Markdown handling
            with open(temp_path, "r", encoding="utf-8") as f:
                text = f.read()

    except Exception as e:
        print(f"Error parsing document: {e}")
        return []
        
    finally:
        # Clean up temp file
        if os.path.exists(temp_path):
            os.remove(temp_path)

    if not text.strip():
        return []

    return [Document(page_content=text, metadata={"source": uploaded_file.name})]