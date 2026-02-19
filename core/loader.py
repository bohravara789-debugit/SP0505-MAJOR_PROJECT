import pdfplumber
import pandas as pd
from docx import Document as DocxDocument
from langchain_core.documents import Document

import pytesseract
from pdf2image import convert_from_bytes


def load_document(uploaded_file):

    uploaded_file.seek(0)
    name = uploaded_file.name.lower()

    text = ""

    if name.endswith(".pdf"):

        # Try normal extraction first
        with pdfplumber.open(uploaded_file) as pdf:
            text = "\n".join([page.extract_text() or "" for page in pdf.pages])

        # If empty → use OCR
        if not text.strip():
            uploaded_file.seek(0)
            images = convert_from_bytes(uploaded_file.read())
            for img in images:
                text += pytesseract.image_to_string(img)

    elif name.endswith(".docx"):
        doc = DocxDocument(uploaded_file)
        text = "\n".join([p.text for p in doc.paragraphs])

    elif name.endswith(".csv"):
        df = pd.read_csv(uploaded_file)
        text = df.to_string()

    else:
        text = uploaded_file.read().decode("utf-8")

    if not text.strip():
        return []

    return [Document(page_content=text)]
